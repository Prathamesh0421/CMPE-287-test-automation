#!/usr/bin/env python3
"""Convert markdown report to DOCX format."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_docx_report(md_file, docx_file):
    """Convert markdown report to formatted DOCX."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Read markdown content
    with open(md_file, 'r') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'Normal'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.5)
                code_lines = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Handle tables
        if line.startswith('|') and '|' in line:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            # Check if next line is not a table line
            if i >= len(lines) or not lines[i].startswith('|'):
                # Process table
                process_table(doc, table_lines)
                in_table = False
                table_lines = []
            continue
        
        # Handle headings
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
        
        # Handle horizontal rules
        elif line.strip() == '---':
            doc.add_paragraph('_' * 80)
        
        # Handle bold text with **
        elif '**' in line:
            p = doc.add_paragraph()
            process_inline_formatting(p, line)
        
        # Handle bullet lists
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        
        # Handle numbered lists
        elif re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s*', '', line)
            p = doc.add_paragraph(text, style='List Number')
        
        # Handle regular paragraphs
        elif line.strip():
            p = doc.add_paragraph(line)
        
        # Empty lines
        else:
            doc.add_paragraph()
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"DOCX report created: {docx_file}")

def process_inline_formatting(paragraph, text):
    """Process inline formatting like bold, italic, code."""
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)

def process_table(doc, table_lines):
    """Process markdown table into DOCX table."""
    # Remove separator line (contains ---)
    table_lines = [line for line in table_lines if not set(line.replace('|', '').replace('-', '').replace(' ', '')) == set()]
    
    if len(table_lines) < 2:
        return
    
    # Parse table
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        rows.append(cells)
    
    # Create table
    if rows:
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Light Grid Accent 1'
        
        # Fill table
        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = cell_text
                # Bold header row
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

if __name__ == "__main__":
    create_docx_report(
        "Weather_Augmentation_Framework_Report.md",
        "Weather_Augmentation_Framework_Report.docx"
    )
